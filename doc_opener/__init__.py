import os
try:
    from variable_test import folder_name, file_name
except ModuleNotFoundError:
    pass
from docx import Document

def open_new_doc(report_folder=os.getcwd()):
    report_folder = os.path.join(os.path.expanduser('~'), 'Desktop', folder_name)
    # Create the Word document path inside the 'report' folder
    file_path = os.path.join(report_folder, file_name)

    # Load the Word document or create a new one if it doesn't exist
    if not os.path.exists(file_path):
        doc = Document()
        doc.save(file_path)

    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        print(paragraph.text)

