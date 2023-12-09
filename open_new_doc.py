import os
from docx import Document
from lxml import etree


def open_new_doc(folder_name, file_name):
    # Create the folder 'report' if it doesn't exist on the user's desktop
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
    report_folder = os.path.join(desktop_path, folder_name)

    if not os.path.exists(report_folder):
        os.makedirs(report_folder)

    # Create the Word document path inside the 'report' folder
    file_path = os.path.join(report_folder, file_name)

    # Load the Word document or create a new one if it doesn't exist
    if not os.path.exists(file_path):
        doc = Document()
        doc.save(file_path)

    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        print(paragraph.text)

# Set the folder name and file name for the Word document
folder_name = 'report'
file_name = 'example.docx'
