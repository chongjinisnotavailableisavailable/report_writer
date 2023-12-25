"""
templates should merge required blocks of text automatically; append in acsending order
list of common pages:
    Main desc
    rationale
    conclusion page
    evidence appendix
    pictures
    disclaimers & citing & etc

"""
from variable_test import file_name, records, variables_dict, client_name,folder_containing_images,generic
from docx import Document
from docx.shared import Inches
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt
from appendix import appendix_1

# a function to determine what kind of merging shall occur to generate a desired template
def template_finder(required_format):
    
    def generic_template():
         required = generic
         return required
    def corporate_template():
        ...
        
    def simplified_template():
        ...
        
    def detailed_template():
        ...
    
    cases = {
        'generic': generic_template,
        'coporate': corporate_template,
        'simplified': simplified_template,
        'detailed': detailed_template        
    }

    required = cases.get(required_format)

    return required()

# A function to bold the document   
# when transplanting variables to merged doc, it should automatically revert it back to normal txt without bold
def bold(document):
    for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    run.bold=True

#inserting objects like tables, logos, pictures in the header section
def insert_header(document):
    for section in document.sections:
        header = section.header

        # for the header portion in the generated template, with the paragraph method, add a run, to add in a picture called logo.jpeg      
        for paragraph in header.paragraphs:
            paragraph.add_run().add_picture('Logo.jpg')

#adds in the TnC/whatever image add the end of the doc.
def TnC(document):
     last_page = document.add_paragraph()

     #on the last page, add_run to insert an image
     #note: the current image size (not file memory size), by default takes up one whole page, and hence a page break is not needed
     run = last_page.add_run()
     image = 'TnC.jpg'

     # this portion forces the image to be of a certain size to take up 1 page, to eliminate the need for a page break
     width_inch = 6.27
     height_inch = 8.76
     run.add_picture(image, width = Inches(width_inch), height = Inches(height_inch))

def insert_tables(document):

    document.add_page_break()

    # p is the title of the table  
    p = document.add_heading('Evidence', level=2) #possibility to use add_heading to create other headers for better organisation
    
    #alignment cant add_run, add_run only works on text/paragraphs
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #treat the newly made header like a paragraph text, and edits it
    # this portion makes it bold, sets colour to black, size to pt 16
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(16) # Adjust the font size as needed
        run.font.color.rgb = RGBColor(0, 0, 0) # RGBColor(0, 0, 0) represents black

    # adds a line break afterwards to make it look nice
    p.add_run().add_break()

    #creates a table of x,y size
    table = document.add_table(rows=1, cols=4)

    #centres the table to the middle of the page
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    #with the header of the cells being called hdr_cells, access the table on the first row with index 0, give it a name with following
    hdr_cells = table.rows[0].cells

    #.text returns a text

    header_texts = ['name', 'size', 'price', 'date']

    #iterates the list to give headers

    # on the first row called headers    
    hdr_cells = table.rows[0].cells

    for index, text in enumerate(header_texts):
        # append left to right index 0 to 3 on header
         hdr_cells[index].text = text

        # bolds, underlines and centralise the header
        # for the cells in the header, bold and underline them       
         header = hdr_cells[index].paragraphs[0]
         formatting = header.runs[0]
         formatting.bold = True
         formatting.underline = True

         # runs or add_runs doesnt work with cells, alightment of this belongs to the cell's property
         # I THINK
         header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #future if statement for conditional remark
    if 'x':
         pass
         ...

    for name, size, price, date in records:
        #iterates over the list and for each list in the list, add a row
        row_cells = table.add_row().cells

        #this portion appends the cell left to right via indexing
        # only accepts type 'str'
        row_cells[0].text = name
        row_cells[1].text = size
        row_cells[2].text = price
        row_cells[3].text = date
    

def insert_key_values(document):
     for paragraph in document.paragraphs:
        for run in paragraph.runs:
             text = run.text
             for key, value in variables_dict.items():
                  text = text.replace(key, value)
             run.text = text

def get_headings_from_filename(file_name):
     parts = file_name.split('_')
     heading_number = parts[0]
     heading_title = parts[1].replace('.docx','').title()
     h = f"{heading_number}.          {heading_title}"
   
     return h

def report_for_who(document):
     p = document.add_paragraph(f'REPORT FOR:          {client_name.upper()}')
     p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

     run = p.runs[0]

     run.bold = True
     run.font.size = Pt(11)
     run.font.name = 'Times New Roman'

     return p

                 
# a function to generate a template base on given input
# default font Times New Roman
def gen_template(files_to_merge, output_folder = os.path.join(os.path.expanduser('~'), 'Desktop','report')):
    
    output_name = file_name
    generated_template = Document()

    report_for_who(generated_template)

    for file_path in files_to_merge:
        doc = Document(file_path)

        font_styles = doc.styles.add_style(name = 'Times New Roman', style_type = 1)
        font = font_styles.font
        font.name = 'Times New Roman'
        
        template_file_name = os.path.basename(file_path)
        heading_title = get_headings_from_filename(template_file_name)

        h = generated_template.add_heading(heading_title, level = 1)
        
        font = h.runs[0].font
        font.name = 'Times New Roman'
        font.bold = True
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)

        h.add_run().add_break()

        for paragraph in doc.paragraphs:
            p = generated_template.add_paragraph()
            p.add_run(paragraph.text).font.name = 'Times New Roman'
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY    
    
    #this section does formatting & adds addtional things from the base template
    #bold(generated_template)  

    insert_header(generated_template)    

    insert_tables(generated_template)

    insert_key_values(generated_template)

    appendix_1(generated_template,folder_containing_images)

    TnC(generated_template)

    #outputs to a path (default desktop/report) with file name (to be set to variable) in variable file
    output_file_path = os.path.join(output_folder, output_name)

    generated_template.save(output_file_path)

    return output_file_path


if __name__ == "__main__":   
    gen_template()