from docx import Document
import os

def insert_objects(generated_template, output_folder = os.path.join(os.path.expanduser('~'), 'Desktop','report')):

    doc = Document(generated_template)

    # for the header portion in the generated template, with the paragraph method, add in a picture called logo.jpeg      
    for section in doc.sections:
        header = section.header
        for document_to_insert in header.paragraphs:
            for run in document_to_insert.runs:
                run.add_picture('Logo.jpg')

    output_file_path = os.path.join(output_folder, 'output_with_logo.docx')            

    doc.save(output_file_path)


    