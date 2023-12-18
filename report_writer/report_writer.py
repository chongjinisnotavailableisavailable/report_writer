import os
from docx import Document
from variable_test import reports_format
from template_generator import gen_template, template_finder

relevant_contents = ["1_Introduction.docx", "2_Description.docx"]

#find or open file with name output in doc_opener
def report_writer_main(files_to_merge = template_finder(reports_format), folder_to_save_in = os.path.join(os.path.expanduser('~'), 'Desktop','report')):
    gen_template(files_to_merge, folder_to_save_in)
    
    """
    use_template()
    insert_headers()
    insert_tables()
    insert_paragraphs()
    replace_keywords()
    insert_pictures
    insert_signor()
    format_report()
    """


## def open_file():
 ##   doc = Document(file_path)

    #iterates over each paragraph, and prints it out in console/env, to be executed in script
   ## for paragraph in doc.paragraphs: 
   ##     print(paragraph.text)

  ##  open_file(file_path)
"""def use_template():
    #path to folder where all the templates are being held
    _path_to_template_folder = os.path.join((os.path.abspath(__file__)),"_templates")

    #get the path where the exact template is
    _path_to_template_file = os.path.join(_path_to_template_folder, reports_format)

    #open template
    source_doc = Document(_path_to_template_file)

    target_doc = Document(file_path)

    #copy template to newly created file in doc_opener
    for template in source_doc.template:
        new_doc = target_doc.add_paragraph(template.text)

    new_doc.save(file_path)

    source_doc.close()
"""





   




if __name__ == "__report_writer__":   
    report_writer_main()