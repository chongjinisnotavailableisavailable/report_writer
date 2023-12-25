from docx.shared import Cm
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import RGBColor, Pt
from variable_test import variables_dict

def split_list_into_chunks(list_to_split):
    chunk_size = 2

    # read as the list to split from i to i + chunk size of 2 (non-inclusive), for that range, split from range(start, stop, step)
    # step is how many times it will move per iteration, namely 2. In other words, how many items i want in the smaller list is determined by this number, chunk_size
    split_list = [list_to_split[i:i+chunk_size] for i in range (0, len(list_to_split), chunk_size)]
    
    return split_list 

def appendix_1(document,folder_containing_images):
    
    # Asks for folder containing images
    folder_path = folder_containing_images

    image_files = [file for file in os.listdir(folder_path) if file.lower().endswith(('.png', '.jpg'))]

    # This portion should loop over the list that has been split into chunks and create page of same formats
    for i in split_list_into_chunks(image_files):
         appendix_organiser(document, folder_path, i)         

def appendix_organiser(document, folder_path, pictures_to_be_added):

    # This portion adds the Appendinx I Heading 
    document.add_page_break()      
    
    appendix = document.add_heading('APPENDIX I', level = 3)
    appendix.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    appendix.add_run().add_break()
    for run in appendix.runs:
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    
    # adds a title text & also replaces it with variable
    title = document.add_paragraph('{$topic}')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        text = run.text
        for key, value in variables_dict.items():
            text = text.replace(key,value)
        run.text = text        
        

    # This portion adds a table to insert pictures
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # This portion is for the image names
    images_name = [] #to be a variable to be entered by user

    # loops through each file name in pictures_to_be_added, retrieves the name part without the extension, and adds it to the images_name list
    for file_name in pictures_to_be_added:
            i = os.path.splitext(file_name)[0]
            images_name.append(i)

    # for the list enumerated in pictures to be added, in an alternating format, add the picture and picture name
    for index, image in enumerate(pictures_to_be_added):
            
        if index %2 == 0:
            cell1 = (index, 0)
            cell2 = (index, 1)
        else:
            cell1 = (index, 1)
            cell2 = (index, 0)

        table.add_row()
        image_cell = table.cell(cell2[0],cell2[1])  # Place each image in a new row
        run = image_cell.paragraphs[0].add_run()
        run.add_picture(os.path.join(folder_path, image), width=Cm(6.0), height=Cm(8.0))
        p = image_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
        name_cell = table.cell(cell1[0],cell1[1])
        run2 = name_cell.paragraphs[0].add_run()
        run2.add_text(images_name[index].upper())
        n = name_cell.paragraphs[0]
        font = n.runs[0].font
        font.name = 'Times New Roman'
        font.bold = True
        font.size = Pt(14)
        font.color.rgb = RGBColor(0, 0, 0)
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
if __name__ == "__main__":   
   appendix_1()